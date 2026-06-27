import csv
import fitz  # PyMuPDF
from email.parser import BytesParser
from email.policy import default
from pathlib import Path
from io import StringIO
from docx import Document as DocxDocument


class TextExtractor:
    """Base class for document text extractors."""

    def extract(self, file_path: str) -> str:
        raise NotImplementedError


class PdfExtractor(TextExtractor):
    """Extract text from PDF files using PyMuPDF."""

    def extract(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        text = "\n\n".join(pages)
        if not text.strip():
            raise ValueError("PDF contains no extractable text (possibly scanned without OCR)")
        return text


class DocxExtractor(TextExtractor):
    """Extract text from .docx Word documents."""

    def extract(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables (flatten to readable format)
        for table in doc.tables:
            text_parts.append("")  # Blank line before table
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
            text_parts.append("")  # Blank line after table

        text = "\n".join(text_parts)
        if not text.strip():
            raise ValueError("DOCX contains no extractable text")
        return text


class EmlExtractor(TextExtractor):
    """Extract text from .eml email files."""

    def extract(self, file_path: str) -> str:
        with open(file_path, "rb") as f:
            msg = BytesParser(policy=default).parse(f)

        text_parts = []

        # Extract headers
        headers_to_include = ["From", "To", "Date", "Subject", "CC", "BCC"]
        for header in headers_to_include:
            if header in msg:
                value = msg[header]
                text_parts.append(f"{header}: {value}")

        text_parts.append("")  # Blank line

        # Extract body (prefer plain text over HTML)
        body = self._extract_body(msg)
        if body:
            text_parts.append(body)

        # List attachments
        attachments = self._list_attachments(msg)
        if attachments:
            text_parts.append("")
            text_parts.append("Attachments:")
            for attachment in attachments:
                text_parts.append(f"  - {attachment}")

        text = "\n".join(text_parts)
        if not text.strip():
            raise ValueError("EML contains no extractable text")
        return text

    def _extract_body(self, msg) -> str:
        """Extract body, preferring plain text over HTML."""
        # Check for multipart message
        if msg.is_multipart():
            for part in msg.iter_parts():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        return payload.decode("utf-8", errors="replace").strip()

            # Fallback to HTML if no plain text found
            for part in msg.iter_parts():
                content_type = part.get_content_type()
                if content_type == "text/html":
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_text = payload.decode("utf-8", errors="replace").strip()
                        return self._html_to_text(html_text)

        # Non-multipart message
        payload = msg.get_payload(decode=True)
        if payload:
            return payload.decode("utf-8", errors="replace").strip()

        return ""

    def _html_to_text(self, html: str) -> str:
        """Simple HTML to text conversion (removes tags)."""
        import re

        text = re.sub(r"<[^>]+>", "", html)
        text = text.replace("&nbsp;", " ")
        text = text.replace("&lt;", "<")
        text = text.replace("&gt;", ">")
        text = text.replace("&amp;", "&")
        return text

    def _list_attachments(self, msg) -> list[str]:
        """List attachment filenames."""
        attachments = []
        if msg.is_multipart():
            for part in msg.iter_parts():
                filename = part.get_filename()
                if filename:
                    attachments.append(filename)
        return attachments


class CsvExtractor(TextExtractor):
    """Extract text from .csv files."""

    def extract(self, file_path: str) -> str:
        text_parts = []

        try:
            # Try to detect encoding
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                if not reader.fieldnames:
                    raise ValueError("CSV has no header row")

                row_count = 0
                for row in reader:
                    if row_count >= 1000:  # Limit to 1000 rows to avoid token explosion
                        text_parts.append(f"... (truncated at 1000 rows)")
                        break

                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                    if row_text.strip():
                        text_parts.append(row_text)
                    row_count += 1

                if row_count == 0:
                    raise ValueError("CSV contains no data rows")

        except UnicodeDecodeError:
            # Fallback to iso-8859-1
            with open(file_path, "r", encoding="iso-8859-1") as f:
                reader = csv.DictReader(f)
                if not reader.fieldnames:
                    raise ValueError("CSV has no header row")

                row_count = 0
                for row in reader:
                    if row_count >= 1000:
                        text_parts.append(f"... (truncated at 1000 rows)")
                        break

                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                    if row_text.strip():
                        text_parts.append(row_text)
                    row_count += 1

                if row_count == 0:
                    raise ValueError("CSV contains no data rows")

        text = "\n".join(text_parts)
        if not text.strip():
            raise ValueError("CSV contains no extractable text")
        return text


class ExtractorFactory:
    """Factory to get appropriate extractor based on file type."""

    _extractors = {
        "pdf": PdfExtractor(),
        "docx": DocxExtractor(),
        "eml": EmlExtractor(),
        "csv": CsvExtractor(),
    }

    @classmethod
    def get_extractor(cls, file_type: str) -> TextExtractor:
        """Get extractor for given file type (lowercase)."""
        extractor = cls._extractors.get(file_type.lower())
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        return extractor


def extract_text_by_type(file_path: str, file_type: str) -> str:
    """
    Extract text from a file based on its type.

    Args:
        file_path: Path to the file
        file_type: File type as string ('pdf', 'docx', 'eml', 'csv')

    Returns:
        Extracted text

    Raises:
        ValueError: If file type not supported or text extraction fails
    """
    extractor = ExtractorFactory.get_extractor(file_type)
    return extractor.extract(file_path)
